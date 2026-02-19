import streamlit as st
import re
import os
import io
import time
import textwrap
import textile
import pymupdf
from fpdf import FPDF
import pandas as pd
from docx import Document
from io import BytesIO
from io import StringIO
from odf.opendocument import OpenDocumentText
from odf.style import Style, TextProperties, ParagraphProperties
from odf.text import H, P
from odf import text, teletype
from odf.opendocument import load

class messages():
    def __init__(self, *args):
        self.label = args[0].strip()
        self.data = args[1]
        self.fileOut = args[2]
        self.mime = args[3]
        self.nFiles = args[4]
        if None not in args:        
            self.mensStr = (f':blue[**{self.fileOut}**] com ***{self.nFiles} arquivo(s)***. Para abri-lo, ' 
                            'pressione (👇) o botão ao lado ➜.') 
            self.mensResult()
            self.exprGer = 'Gerado o arquivo para download❗'
        else:
            self.exprGer = 'Gerado o conteúdo para exibição em tela❗'
        
    def mensResult(self):
        nLabel = len(self.label)
        nMensStr = round(len(self.mensStr)/3, 0)
        nStr = nLabel + nMensStr
        if self.label == 'ODT':
            y = 3.1
        elif self.label == 'HTML':
            y = 3.5
        else:
            if nStr == 38:
                y = 3.1
            elif nStr == 39:
                y = 3.4
            elif nStr == 40:
                y = 3.8
            else:
                y = 4.0
        colMens, colDown = st.columns([nMensStr, y], width='stretch', vertical_alignment='center')
        colDown.download_button(
            label=self.label,
            data=self.data,
            file_name=self.fileOut,
            mime=self.mime,
            icon=':material/download:'
        )
        colMens.success(self.mensStr, icon='✔️',  width='stretch')
        self.label = f'"{self.mensStr}".'
        self.mensExib()
        
    def mensToast(self): 
        textToast = ['1️⃣ Para maiores detalhes, consulte a opção Detalhes do app.', 
                    '2️⃣ Nela, você encontrará Formatos do app e Funcionalidades do app.', 
                    ('3️⃣ Devido a problemas de formatação, o texto resultante poderá conter ' 
                     'símbolos estranhos.'),   
                    '4️⃣ É sempre recomendável a conferência com o original.',
                    '5️⃣ O arquivo convertido não conservará nem herdará a formatação primitiva.',
                    '6️⃣ Selecionado arquivo PDF, verificar se é pesquisável ou tem OCR.']
        msg = st.toast('🪄 Espere a exibição destas :violet[**6 dicas fundamentais**]❗')
        textToast.insert(0, '(✋ Evite fechar as janelas (⿻)❗)')
        for text in textToast: 
            time.sleep(2.5)
            msg.toast(text)
        
    @st.dialog('⚠️ Falha no app❗')
    def mensError(self, str):
        st.markdown(f'{str} Entre em contato com o administrador da ferramenta!')
        
    @st.dialog(f'✅️ Resultado bem-sucedido!')
    def mensExib(self):        
        st.markdown(f'Role o :violet[**mouse**] (🖱) para o contêiner '
                    f'localizado :violet[**abaixo**] (￬) de "***📌 Detalhes do app***" '
                    f'e encontre ***{self.label}***')

class operatorsFiles():
    def __init__(self, *args):
        self.nFiles = args[0]
        
    @st.cache_data
    def docxToTxt(_self, fileDown):
        doc = Document(fileDown)
        textAll = [para.text + '\n' for para in doc.paragraphs]
        return textAll
        
    @st.cache_data
    def pdfToTxt(_self, fileDown):
        with open('prov.pdf', 'wb') as f:
            f.write(fileDown.getbuffer())
        pages = pymupdf.open('prov.pdf')
        textAll = ''
        for page in pages:
            textAll += page.get_text() + '\n'
        return textAll
        
    @st.cache_data
    def odtToTxt(_self, fileDown):
        doc = load(fileDown)
        allParag = doc.getElementsByType(text.P)
        textAll = [teletype.extractText(parag) for parag in allParag]
        return textAll
        
    @st.cache_data
    def rtfToTxt(_self, fileDown):
        textAll = fileDown.read().decode('utf-8')
        return textAll
        
    def txtToScroll(self, textFileAll, allNames): 
        self.line = '—'*75
        nFiles = len(textFileAll)
        if nFiles <= 1:
            strFile = 'do único arquivo selecionado.'
        else:
            strFile = f'dos {nFiles} arquivos selecionados.'
        textFinal = f'🖍️ ***Nome, numeração e conteúdo {strFile}***<br>'
        textIni = f'"{textFinal}".'
        nLines = 0
        for t, textFile in enumerate(textFileAll):
            if type(textFile) is list:
                textFile = ''.join(textFile)
            nLines += len(textFile.split('\n'))
            nameFile = f'🗁 {allNames[t]}   ➤  {t+1} de {nFiles}'
            prefix = f'{self.line}<br>{nameFile}<br>'
            content = f'{self.line}<br>{textFile}<br>'   
            content = content.replace('\n', '<br>')
            textFinal += f'<br>{prefix}{content}<br>'
        hgtLines = 30 * (nLines + 4)
        if hgtLines >= 2500:
            hgtLines = 2500
        with st.container(height=hgtLines):
            st.markdown(textFinal, unsafe_allow_html=True)
        objMens = messages(textIni, None, None, None, None)
        objMens.mensExib()
        
    def txtToTxt(self, textFile):
        self.textFile = textFile
        self.mode = 1
        self.depureText()
        self.fileOut = 'arquivo_resultante.txt'
        try:
            self.textPure = self.textPure.encode('cp1252').decode('utf-8')
        except:
            try:
                self.textPure = self.textPure.encode('latin-1').decode('latin-1')
            except: 
                self.textPure = self.textPure.encode('utf-8').decode('utf-8')            
        messages('TXT', self.textPure, self.fileOut, 'text/plain', self.nFiles)
        
    def txtToDocx(self, textFile):
        self.textFile = textFile
        self.mode = 1
        self.depureText()
        self.fileOut = 'arquivo_resultante.docx' 
        self.createDocx()
        messages('DOCX', self.ioDocx, self.fileOut, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                 self.nFiles)
    
    def createDocx(self):
        doc = Document()
        doc.add_paragraph(self.textPure)
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        self.ioDocx = bio.getvalue()
        
    def txtToRtf(self, textFile):
        self.textFile = textFile
        self.mode = 1
        self.depureText()
        self.fileOut = 'arquivo_resultante.rtf'
        try:
            self.textPure = self.textPure.encode('cp1252').decode('utf-8')
        except:
            try:
                self.textPure = self.textPure.encode('latin-1').decode('latin-1')
            except: 
                self.textPure = self.textPure.encode('utf-8').decode('utf-8') 
        messages('RTF', self.textPure, self.fileOut, 'application/rtf', self.nFiles)
    
    def txtToHtml(self, textFile):
        self.textFile = textFile
        self.mode = 0
        self.depureText()
        self.fileOut = 'arquivo_resultante.html'
        messages('HTML', self.textPure, self.fileOut, 'text/html', self.nFiles)
    
    def txtToOdt(self, textFile): 
        self.textFile = textFile
        self.mode = 1
        self.depureText()
        self.fileOut = 'arquivo_resultante.odt' 
        self.createOdt()
        messages('ODT', self.ioOds, self.fileOut, 'application/vnd.oasis.opendocument.text', 
                 self.nFiles)
                 
    def txtToXhtml(self, textFile):
        self.textFile = textFile
        self.mode = 0
        self.depureText()
        xhtmlContent = textile.textile(self.textPure, html_type='xhtml')
        self.fileOut = 'arquivo_resultante.xhtml'
        messages('XHTML', xhtmlContent, self.fileOut, 'application/xhtml+xml', self.nFiles)
    
    def txtToPdf(self, textFile):
        self.textFile = textFile
        self.mode = 1
        self.depureText()
        pdf = FPDF('L', 'mm', 'A4')   
        pdf.add_page()
        pdf.set_font("arial", size = 12)
        textSplit = self.textPure.split('\n')
        for text in textSplit: 
            lines = textwrap.wrap(text, width=145)
            for line in lines:
                try:
                    pdf.cell(300, 10, txt = line, ln = 1, align = 'L')
                except:
                    pass
        pdf.output('arquivo_resultante.pdf')
        with open('arquivo_resultante.pdf', "rb") as f:
            pdfBytes = f.read()
        self.fileOut = 'arquivo_resultante.pdf'
        messages('PDF', pdfBytes, self.fileOut, 'application/pdf', self.nFiles)
    
    def createOdt(self):
        doc = OpenDocumentText() 
        parag = P()
        teletype.addTextToElement(parag, self.textPure)
        doc.text.addElement(parag)
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        self.ioOds = bio.getvalue()
        
    def depureText(self):
        self.textPure = ''
        for text in self.textFile:
            if type(text) is list:
                text = ''.join(text)
            if self.mode == 0:
                self.textPure += '<br>' + text.replace('\n', '<br>')
            else:
                self.textPure += text + '\n'
        
class main():
    def __init__(self): 
        self.setPage()
        self.exts = sorted(['txt', 'csv', 'tsv', 'py', 'json', 'js', 'html', 'docx', 'odt', 
                            'rtf', 'log', 'bat', 'php', 'css', 'xml', 'msg', 'md', 'cfg', 'conf', 
                            'jsp', 'cpp', 'sql', 'jspx', 'tex', 'xhtml', 'pdf'])
        self.exts = [ext.upper() for ext in self.exts]
        self.nExts = len(self.exts)
        self.extsStr = ', '.join(self.exts[:-1])
        self.extsStr += f' e {self.exts[-1]}'
        self.extOrders = ['HTML', 'TXT', 'DOCX', 'RTF', 'ODT', 'XHTML', 'PDF']
        self.locateMedia() 
        self.buttons = {0: ['TELA', 'screenOnly', ':material/screen_search_desktop:', 
                            'Exibe o conteúdo dos arquivos na tela.'], 
                        1: [self.extOrders[1], 'contentTxt', ':material/text_ad:',
                            'Transforma o conteúdo em arquivo TXT.'],
                        2: [self.extOrders[0], 'contentHtml', ':material/folder_code:', 
                            'Transforma o conteúdo em arquivo HTML.'], 
                        3: [self.extOrders[2], 'contentDocx', ':material/docs:', 
                            'Transforma o conteúdo em arquivo DOCX.'], 
                        4: [self.extOrders[3], 'contentRtf', ':material/note_alt:', 
                            'Transforma o conteúdo em arquivo RTF.'], 
                        5: [self.extOrders[4], 'contentOdf', ':material/content_paste:', 
                            'Transforma o conteúdo em arquivo ODF.'],
                        6: [self.extOrders[5], 'contentXhtml', ':material/code_blocks:', 
                            'Transforma o conteúdo em arquivo XHTML.'], 
                        7: [self.extOrders[6], 'contentPdf', ':material/picture_in_picture_center:', 
                            'Transforma o conteúdo em arquivo PDF.']}
        self.values = [self.buttons[b] for b in range(len(self.buttons))]
        self.nameOpts = [value[0] for value in self.values]
        self.keyButts = [value[1] for value in self.values]
        self.disabs = [True for w in range(len(self.keyButts))]
        with st.container(border=4):
            colDown, colButton = st.columns([25, 17], width='stretch')
            with colDown:
                with st.container(border=None):
                    st.space(size='small')
                    self.options = st.multiselect(label=f'🏷️ Formatos para seleção exclusiva', 
                                                  options=self.exts, 
                                                  placeholder='Únicos formatos desejados',  
                                                  help='Restringe os tipos de arquivo desejados. '
                                                       f'Se nada for escolhido, valerão estes :blue[***{self.nExts} formatos***]:\n'
                                                       f'{self.extsStr}.', 
                                                  disabled=st.session_state.multSel)
                    if self.options == []:
                        typeSel = self.exts
                        helpStr = (f'Selecione ou arraste arquivos com qualquer um destes '
                                   f':blue[***{self.nExts} formatos***]:\n{self.extsStr}.')
                    else:
                        typeSel = self.options
                        if len(self.options) == 1:
                            helpStr = f'Selecione ou arraste arquivos com formato :blue[***{self.options[0]}***].'
                        else:
                            self.nOptions = len(self.options)
                            self.optionsStr = ', '.join(self.options[:-1])
                            self.optionsStr += f' e {self.options[-1]}'
                            helpStr = (f'Selecione ou arraste arquivos com qualquer um destes '
                            f':blue[***{self.nOptions} formatos***]:\n{self.optionsStr}.')
                    st.space(size='small')
                    self.upDowns = st.file_uploader(label='📚 Seleção ou arrastamento de arquivos', 
                                                    accept_multiple_files=True, 
                                                    type=typeSel, 
                                                    help=helpStr, 
                                                    max_upload_size=1024*20, 
                                                    width='stretch', 
                                                    disabled=st.session_state.fileDown)
                    self.nDowns = len(self.upDowns)
                    self.nKeys = len(self.keyButts)
                    if self.nDowns:
                        self.disabs = [False for w in range(self.nKeys)]
                        if self.nDowns == 1:
                            self.checkExt()
                    else:
                        self.disabs = [True for w in range(self.nKeys)]
            with colButton:
                with st.container(border=None):
                    if self.nDowns == 1:
                        st.space(size='xxsmall')
                    elif self.nDowns == 2:
                        st.space(size='xxsmall')
                        st.space(size='xxsmall')
                    elif self.nDowns > 2:
                        st.space(size='medium')
                    colIso, colAll = st.columns(spec=2, width='stretch')
                    self.screen = colIso.button(self.buttons[0][0], key=self.buttons[0][1], icon=self.buttons[0][2], 
                                                help=self.buttons[0][3], use_container_width=True, 
                                                disabled=self.disabs[0])
                    self.txt = colAll.button(self.buttons[1][0], key=self.buttons[1][1], icon=self.buttons[1][2], 
                                             help=self.buttons[1][3], use_container_width=True, 
                                             disabled=self.disabs[1])
                    colHtml, colDocx = st.columns(spec=2, width='stretch')
                    self.html = colHtml.button(self.buttons[2][0], key=self.buttons[2][1], icon=self.buttons[2][2], 
                                               help=self.buttons[2][3], use_container_width=True, 
                                               disabled=self.disabs[2])
                    self.docx = colDocx.button(self.buttons[3][0], key=self.buttons[3][1], icon=self.buttons[3][2], 
                                               help=self.buttons[3][3], use_container_width=True, 
                                               disabled=self.disabs[3])
                    colRtf, colOds = st.columns(spec=2, width='stretch')
                    self.rtf = colRtf.button(self.buttons[4][0], key=self.buttons[4][1], icon=self.buttons[4][2], 
                                               help=self.buttons[4][3], use_container_width=True, 
                                               disabled=self.disabs[4])
                    self.ods = colOds.button(self.buttons[5][0], key=self.buttons[5][1], icon=self.buttons[5][2], 
                                               help=self.buttons[5][3], use_container_width=True, 
                                               disabled=self.disabs[5]) 
                    colXhtml, colPdf = st.columns(spec=2, width='stretch')
                    self.xhtml = colXhtml.button(self.buttons[6][0], key=self.buttons[6][1], icon=self.buttons[6][2], 
                                               help=self.buttons[6][3], use_container_width=True, 
                                               disabled=self.disabs[6])
                    self.pdf = colPdf.button(self.buttons[7][0], key=self.buttons[7][1], icon=self.buttons[7][2], 
                                               help=self.buttons[7][3], use_container_width=True, 
                                               disabled=self.disabs[7])
                    if self.nDowns == 2:
                        st.space(size='xxsmall')       
        with st.expander(label='Detalhes do app', 
                         expanded=False,
                         icon='📌', 
                         width='stretch'):
            tabOne, tabTwo = st.tabs(['Formatos do app', 
                                      'Funcionalidades do app'], 
                                      width='stretch')
            with tabOne:
                self.formatTab(0)
            with tabTwo:
                self.formatTab(1)
        if not st.session_state.toast:
            objMens = messages('toast', None, None, None, None)
            objMens.mensToast() 
            st.session_state.toast = True
            st.session_state.multSel = False
            st.session_state.fileDown = False
            st.rerun()
        if self.nDowns >= 1:
            self.processDown() 
    
    def checkExt(self):
        extDowns = [os.path.splitext(down.name)[1].replace('.', '').upper() for down in self.upDowns]
        try:
            indDown = self.extOrders.index(extDowns[0])
            if indDown == 0:
                indDown = 2
            elif indDown != 1:
                indDown += 1
            self.disabs[indDown] = True
        except:
            pass
            
    def formatTab(self, mode):
        if mode == 0:
            self.formatExpander()
            optData = pd.DataFrame(self.expandFiles, index = self.index)
            st.table(optData)
        elif mode == 1:
            tabScreen, tabTxt, tabHtml, tabDocx, tabRtf, tabOdt, tabXhtml, tabPdf = st.tabs(self.nameOpts,
                                                                                            width='stretch', 
                                                                                            default=None)
            self.opts = ['Informações?  ℹ', 
                         'Imagens?  📸', 
                         'Vídeos?  :movie_camera:']
            textScreen = ('1️⃣ Exibe na tela, em formato :violet[**texto**] (📝) e logo '
                          ':violet[**abaixo**] (⇩) deste contêiner, o conteúdo dos arquivos '
                          'selecionados.\n\n'
                          '2️⃣  É a única opção em que o conteúdo aparece :violet[**segmentado**] (✂️), '
                          'isto é, arquivo por arquivo.\n\n')
            textLbd = lambda a: (f'1️⃣ Gera botão para :violet[**download**] (📥) de arquivo :violet[**{a}**] único com '
                                 'o conteúdo dos arquivos selecionados.\n\n'
                                 f'2️⃣ Se o único arquivo selecionado tiver a extensão :violet[**.{a.lower()}**], o '
                                 ':violet[**botão**] (◻️) correspondente será :violet[**desativado**] (🚫).')
            textIncr = list(map(textLbd, self.extOrders))
            textAdds = (f'3️⃣ Devido a problemas de :violet[**formatação**] (⚙️), '
                        'o texto resultante poderá conter símbolos :violet[**estranhos**] (�).\n\n'  
                        '4️⃣ É sempre recomendável a :violet[**conferência**] (↔) com o original.\n\n'
                        '5️⃣ O arquivo convertido não conservará nem herdará a formatação primitiva.\n\n'
                        '6️⃣ Se qualquer dos arquivos selecionados for :violet[**PDF**], a :violet[**extração**] '
                        'de texto dependerá de ser ele :violet[**pesquisável**] (🔎) na origem ou após aplicação'
                        'de :violet[**OCR**] (https://pt.wikipedia.org/wiki/Reconhecimento_%C3%B3tico_de_caracteres).')
            lambImg = lambda a: [file for file in self.dictMedia[self.keysMedia[a]] if file.find('image') >= 0]
            lambVd = lambda a: [file for file in self.dictMedia[self.keysMedia[a]] if file.find('vídeo') >= 0]
            self.imgVd = {0:[lambImg(-1), lambVd(-1), f'{textScreen}\n\n{textAdds}'], 
                          1:[lambImg(0), lambVd(0), f'{textIncr[0]}\n\n{textAdds}'], 
                          2:[lambImg(1), lambVd(1), f'{textIncr[1]}\n\n{textAdds}'], 
                          3:[lambImg(2), lambVd(2), f'{textIncr[2]}\n\n{textAdds}'],
                          4:[lambImg(3), lambVd(3), f'{textIncr[3]}\n\n{textAdds}'], 
                          5:[lambImg(4), lambVd(4), f'{textIncr[4]}\n\n{textAdds}'],
                          6:[lambImg(5), lambVd(5), f'{textIncr[5]}\n\n{textAdds}'], 
                          7:[lambImg(6), lambVd(6), f'{textIncr[6]}\n\n{textAdds}']}
            with tabScreen:
                self.num = 0
                self.imgVdSel = self.imgVd[self.num]
                self.formatImgVd()
            with tabHtml:
                self.num = 1
                self.imgVdSel = self.imgVd[self.num]
                self.formatImgVd()
            with tabTxt:
                self.num = 2
                self.imgVdSel = self.imgVd[self.num]
                self.formatImgVd()
            with tabDocx:
                self.num = 3
                self.imgVdSel = self.imgVd[self.num]
                self.formatImgVd()
            with tabRtf:
                self.num = 4
                self.imgVdSel = self.imgVd[self.num]
                self.formatImgVd()
            with tabOdt:
                self.num = 5
                self.imgVdSel = self.imgVd[self.num]
                self.formatImgVd()  
            with tabXhtml:
                self.num = 6
                self.imgVdSel = self.imgVd[self.num]
                self.formatImgVd()
            with tabPdf:
                self.num = 7
                self.imgVdSel = self.imgVd[self.num]
                self.formatImgVd() 
            
    def locateMedia(self):
        self.dictMedia = {ext:[] for ext in self.extOrders}
        self.dictMedia['screen'] = []
        self.keysMedia = list(self.dictMedia.keys())
        dirMedia = 'media/'
        files = [file for file in os.listdir(dirMedia) if os.path.splitext(file)[1] == '.jpg']
        files += [file for file in os.listdir(dirMedia) if os.path.splitext(file)[1] == '.webm']
        for key in self.keysMedia:
            for file in files:
                if file.split('_')[1].strip() == key: 
                    self.dictMedia[key].append(os.path.join(dirMedia, file))
    
    def formatImgVd(self):
        with st.container(border=4):
            opt = st.radio('***Escolha imagem ou vídeo da funcionalidade.***', 
                           self.opts, horizontal=True, index=None, width='stretch', 
                           key=f'radio{self.num}')
            if opt == self.opts[0]:
                try:
                    st.info(self.imgVd[self.num][-1], width='stretch')                  
                except:
                    pass
            elif opt == self.opts[1]:
                imgs = self.imgVdSel[0]
                nImgs = len(imgs)
                if nImgs == 1:
                    exprImg = f'{nImgs} imagem.'
                else:
                    exprImg = f'{nImgs} imagens.'
                posImg = st.slider(label=f'Deslize o botão para escolher umas das {exprImg}', 
                                   min_value=1, max_value=nImgs, value=1, key=f'key_{self.num}img')
                st.image(imgs[posImg-1], width=600)
            elif opt == self.opts[2]:
                vds = self.imgVdSel[1]
                nVds = len(vds)
                if nVds == 1:
                    exprVds = f'{nVds} vídeo.'
                else:
                    exprVds = f'{nVds} vídeos.'
                posVds = st.slider(label=f'Deslize o botão para escolher um dos {exprVds}', 
                                min_value=1, max_value=nVds, value=1, key=f'key_{self.num}vds')
                st.video(vds[posVds-1], autoplay=True, width='stretch')
    
    def formatExpander(self):
        unionExts = list(set(self.extOrders + self.exts))
        commonExts = sorted(list(set(self.exts) & set(self.extOrders)))
        filesExt = {'TXT': 'https://pt.wikipedia.org/wiki/Arquivo_de_texto', 
                    'DOCX': 'https://www.onlyoffice.com/blog/pt-br/2024/03/docx', 
                    'HTML': 'https://pt.wikipedia.org/wiki/HTML', 
                    'RTF': 'https://pt.wikipedia.org/wiki/Rich_Text_Format', 
                    'ODT': 'https://pt.wikipedia.org/wiki/OpenDocument', 
                    'PDF': 'https://pt.wikipedia.org/wiki/PDF', 
                    'XHTML': 'https://pt.wikipedia.org/wiki/XHTML', 
                    'BAT': 'https://pt.wikipedia.org/wiki/Batch', 
                    'CFG': 'https://en.wikipedia.org/wiki/Configuration_file', 
                    'CONF': 'https://en.wikipedia.org/wiki/Configuration_file', 
                    'CPP': 'https://ficheiros.com.br/extensao/cpp/', 
                    'CSS': 'https://pt.wikipedia.org/wiki/CSS', 
                    'CSV': 'https://pt.wikipedia.org/wiki/Comma-separated_values', 
                    'JS': 'https://pt.wikipedia.org/wiki/JavaScript', 
                    'JSON': 'https://pt.wikipedia.org/wiki/JSON', 
                    'JSP': 'https://pt.wikipedia.org/wiki/JavaServer_Pages', 
                    'JSPX': 'https://pt.wikipedia.org/wiki/JavaServer_Pages', 
                    'LOG': 'https://aws.amazon.com/pt/what-is/log-files/', 
                    'MD': 'https://pt.wikipedia.org/wiki/Markdown', 
                    'MSG': 'https://www.geeksforgeeks.org/techtips/msg-text-format/', 
                    'PHP': 'https://pt.wikipedia.org/wiki/PHP', 
                    'PY': 'https://en.wikipedia.org/wiki/Python_(programming_language)', 
                    'SQL': 'https://pt.wikipedia.org/wiki/SQL', 
                    'TEX': 'https://pt.wikipedia.org/wiki/TeX', 
                    'TSV': 'https://en.wikipedia.org/wiki/Tab-separated_values', 
                    'XHTML': 'https://pt.wikipedia.org/wiki/XHTML', 
                    'XML': 'https://pt.wikipedia.org/wiki/XML'}
        self.allExts = sorted(list(filesExt.keys()))
        statusExt = []
        for ext in self.allExts:
            testOne = ext in commonExts
            testTwo = all([ext in self.exts, ext not in self.extOrders])
            testThree = all([ext in self.extOrders, ext not in self.exts])
            if testOne: 
                status = '✦ seleção (📂) ➕ conversão (🔄)'
            elif testTwo:
                status = '🔸 somente conversão (🔄)'
            elif testThree:
                status = '🔸 somente seleção (📂)'
            statusExt.append(status)
        listUrls = [filesExt[ext] for ext in self.allExts]
        self.expandFiles = {'Tipo de arquivo': self.allExts, 
                            'URL': listUrls, 
                            'Função': statusExt}
        self.index = [w + 1 for w in range(len(self.allExts))]   
    
    def processDown(self):
        try:
            self.textFileAll = []
            self.namesAll = []
            objOperat = operatorsFiles(self.nDowns)
            for d, down in enumerate(self.upDowns):
                nameDown = down.name
                keyDown = f'{nameDown}_{d+1}'
                getDown = down.getvalue()
                try:
                    stringIO = StringIO(getDown.decode('latin-1'))
                except:
                    stringIO = StringIO(getDown.decode('cp1252'))
                textFile = stringIO.read()
                ext = os.path.splitext(nameDown)[1]
                if ext == '.docx':
                    textFile = objOperat.docxToTxt(down)
                elif ext.lower() == '.pdf':
                    textFile = objOperat.pdfToTxt(down)
                elif ext == '.odt':
                    textFile = objOperat.odtToTxt(down)
                elif ext == '.rtf':
                    textFile = objOperat.rtfToTxt(down)
                self.textFileAll.append(textFile)
                self.namesAll.append(nameDown)
            with st.spinner('❯❯❯❯ Operação em andamento...', show_time=True, width='content'):
                if self.screen: 
                    objOperat.txtToScroll(self.textFileAll, self.namesAll)
                elif self.txt:
                    objOperat.txtToTxt(self.textFileAll)
                elif self.docx:
                    objOperat.txtToDocx(self.textFileAll)
                elif self.html: 
                    objOperat.txtToHtml(self.textFileAll)
                elif self.rtf:
                    objOperat.txtToRtf(self.textFileAll)
                elif self.ods:
                    objOperat.txtToOdt(self.textFileAll)
                elif self.xhtml:
                    objOperat.txtToXhtml(self.textFileAll)
                elif self.pdf:
                    objOperat.txtToPdf(self.textFileAll)
        except Exception as error: 
            textError = f'Ocorreu o seguinte erro na operação:\n:green[***{error}***].' 
            objMens = messages(textError, None, None, None, None)
            objMens.mensError(textError) 
        
    def setPage(self):
        st.set_page_config(
        page_title="Ex-stream-ly Cool App",
        page_icon="🧊",
        layout="wide")   
        with open('configCss.css') as f:
            css = f.read()
        st.markdown(f'<style>{css}</style>', unsafe_allow_html=True)

if __name__ == '__main__':
    if 'toast' not in st.session_state:
        st.session_state.toast = False
    if 'multSel' not in st.session_state:
        st.session_state.multSel = True
    if 'fileDown' not in st.session_state:
        st.session_state.fileDown = True
    main()



